from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas




def build_docx_from_text(text: str) -> BytesIO:
    buffer = BytesIO()
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(buffer)
    buffer.seek(0)
    return buffer


from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm


def build_pdf_from_text(text: str) -> BytesIO:
    buffer = BytesIO()

    # Définition de la page A4
    pdf = canvas.Canvas(buffer, pagesize=A4)

    width, height = A4

    # Marges
    left_margin = 2 * cm
    right_margin = width - 2 * cm
    top_margin = height - 2 * cm
    bottom_margin = 2 * cm

    # Police
    pdf.setFont("Helvetica", 12)

    # Position de départ du texte
    y = top_margin

    # Largeur maximum du texte
    max_width = right_margin - left_margin

    # Découper le texte automatiquement en lignes qui tiennent sur la page
    from reportlab.lib.utils import simpleSplit
    lines = simpleSplit(text, "Helvetica", 12, max_width)

    for line in lines:
        if y < bottom_margin:     # Nouvelle page si plus de place
            pdf.showPage()
            pdf.setFont("Helvetica", 12)
            y = top_margin

        pdf.drawString(left_margin, y, line)
        y -= 15  # Espacement entre lignes

    pdf.save()
    buffer.seek(0)
    return buffer

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from datetime import datetime
from reportlab.lib.utils import simpleSplit


def build_pdf_lm(nom, prenom, adresse, email, tel, destinataire, objet, contenu):
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)

    width, height = A4
    left = 2.2 * cm
    right = width - 2.2 * cm
    top = height - 2.5 * cm
    bottom = 2 * cm

    y = top

    # Coordonnées du candidat
    pdf.setFont("Helvetica", 11)
    header = [
        f"{prenom} {nom}",
        adresse,
        email,
        tel,
        "",
        datetime.now().strftime("%d %B %Y"),
        "",
    ]

    for h in header:
        pdf.drawString(left, y, h)
        y -= 15

    # Destinataire
    if destinataire.strip():
        for line in destinataire.split("\n"):
            pdf.drawString(left, y, line)
            y -= 15
        y -= 10

    # Objet
    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawString(left, y, f"Objet : {objet}")
    y -= 25
    pdf.setFont("Helvetica", 11)

    # Corps du texte automatique
    max_width = right - left
    lines = simpleSplit(contenu, "Helvetica", 11, max_width)

    for line in lines:
        if y < bottom:
            pdf.showPage()
            y = top
            pdf.setFont("Helvetica", 11)

        pdf.drawString(left, y, line)
        y -= 15

    # Signature
    y -= 30
    pdf.drawString(left, y, "Cordialement,")
    y -= 25
    pdf.drawString(left, y, f"{prenom} {nom}")

    pdf.save()
    buffer.seek(0)
    return buffer

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.colors import Color, HexColor
from reportlab.lib.units import cm
from reportlab.lib.utils import simpleSplit
from io import BytesIO


def build_pdf_cv_modern(cv_data):
    """
    cv_data doit contenir :
    {
      "nom": "...",
      "prenom": "...",
      "email": "...",
      "tel": "...",
      "ville": "...",
      "formation": "...",
      "experiences": "...",
      "projets": "...",
      "tech": "...",
      "soft": "...",
      "interets": "...",
    }
    """

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    # -------- COULEURS --------
    bleu = HexColor("#003A63")
    bleu_clair = HexColor("#2B6CA3")
    gris = HexColor("#444444")

    # -------- LAYOUT --------
    left_col_width = 6 * cm
    margin = 1.5 * cm

    # -------- BACKGROUND GAUCHE --------
    pdf.setFillColor(bleu)
    pdf.rect(0, 0, left_col_width, height, fill=True, stroke=False)

    # -------- NOM & PRÉNOM --------
    pdf.setFillColor("white")
    pdf.setFont("Helvetica-Bold", 18)
    pdf.drawString(margin, height - 2*cm, f"{cv_data['prenom']} {cv_data['nom']}")

    # -------- CONTACT --------
    pdf.setFont("Helvetica", 11)
    y = height - 3.5*cm
    for line in [
        cv_data["email"],
        cv_data["tel"],
        cv_data["ville"]
    ]:
        pdf.drawString(margin, y, line)
        y -= 14

    # -------- TITRE À DROITE --------
    pdf.setFillColor(gris)
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(left_col_width + margin, height - 2*cm, "CV Étudiant – Profil Parcoursup")

    # -------- SECTION FORMATION --------
    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawString(left_col_width + margin, height - 3.5*cm, "🎓 Formation")

    pdf.setFont("Helvetica", 11)
    text = simpleSplit(cv_data["formation"], "Helvetica", 11, width - left_col_width - 2*margin)
    y = height - 4.5*cm
    for line in text:
        pdf.drawString(left_col_width + margin, y, line)
        y -= 14

    # -------- SECTION EXPÉRIENCES --------
    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawString(left_col_width + margin, y - 10, "💼 Expériences")
    y -= 30

    pdf.setFont("Helvetica", 11)
    text = simpleSplit(cv_data["experiences"], "Helvetica", 11, width - left_col_width - 2*margin)
    for line in text:
        pdf.drawString(left_col_width + margin, y, line)
        y -= 14

    # -------- SECTION PROJETS --------
    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawString(left_col_width + margin, y - 10, "🚀 Projets")
    y -= 30

    pdf.setFont("Helvetica", 11)
    text = simpleSplit(cv_data["projets"], "Helvetica", 11, width - left_col_width - 2*margin)
    for line in text:
        pdf.drawString(left_col_width + margin, y, line)
        y -= 14

    # -------- SECTION COMPÉTENCES (COLONNE GAUCHE) --------
    pdf.setFillColor("white")
    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawString(margin, height/2, "⭐ Compétences techniques")

    pdf.setFont("Helvetica", 10)
    y2 = height/2 - 20
    tech_lines = simpleSplit(cv_data["tech"], "Helvetica", 10, left_col_width - 2*margin)
    for line in tech_lines:
        pdf.drawString(margin, y2, line)
        y2 -= 12

    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawString(margin, y2 - 20, "🤝 Compétences humaines")

    pdf.setFont("Helvetica", 10)
    y2 -= 40
    soft_lines = simpleSplit(cv_data["soft"], "Helvetica", 10, left_col_width - 2*margin)
    for line in soft_lines:
        pdf.drawString(margin, y2, line)
        y2 -= 12

    # -------- INTÉRÊTS --------
    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawString(margin, y2 - 25, "❤️ Centres d'intérêt")

    pdf.setFont("Helvetica", 10)
    y2 -= 45
    inter_lines = simpleSplit(cv_data["interets"], "Helvetica", 10, left_col_width - 2*margin)
    for line in inter_lines:
        pdf.drawString(margin, y2, line)
        y2 -= 12

    pdf.save()
    buffer.seek(0)
    return buffer

def build_pdf_lm(prenom_nom, adresse, email, tel, destinataire, objet, contenu):
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)

    width, height = A4
    left = 2.2 * cm
    right = width - 2.2 * cm
    top = height - 2.5 * cm
    bottom = 2 * cm

    y = top

    # Coordonnées du candidat
    pdf.setFont("Helvetica", 11)
    header = [
        prenom_nom,
        adresse,
        email,
        tel,
        "",
        datetime.now().strftime("%d %B %Y"),
        "",
    ]

    for h in header:
        pdf.drawString(left, y, h)
        y -= 15

    # Destinataire
    if destinataire.strip():
        for line in destinataire.split("\n"):
            pdf.drawString(left, y, line)
            y -= 15
        y -= 10

    # Objet
    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawString(left, y, f"Objet : {objet}")
    y -= 25
    pdf.setFont("Helvetica", 11)

    # Corps du texte automatique
    max_width = right - left
    lines = simpleSplit(contenu, "Helvetica", 11, max_width)

    for line in lines:
        if y < bottom:
            pdf.showPage()
            y = top
            pdf.setFont("Helvetica", 11)

        pdf.drawString(left, y, line)
        y -= 15

    # Signature
    y -= 30
    pdf.drawString(left, y, "Cordialement,")
    y -= 25
    pdf.drawString(left, y, prenom_nom)

    pdf.save()
    buffer.seek(0)
    return buffer

    
from io import BytesIO
from docx import Document
from docx.shared import Cm, Pt

def build_docx_from_text(text: str) -> BytesIO:
    buffer = BytesIO()
    doc = Document()

    # --- MARGES PRO ---
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # --- STYLE DE BASE ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    # --- AJOUT DES PARAGRAPHES ---
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]

    for p in paragraphs:
        doc.add_paragraph(p)

    doc.save(buffer)
    buffer.seek(0)
    return buffer
